from docx import Document
import os


SAMPLE_DOCUMENT = ""
DEFAULT_FONT_PT_SIZE = 11


try:
    import local_settings
except ImportError:
    SAMPLE_DOCUMENT = os.path.join("documents", "doc_templates", "sample.docx")
else:
    SAMPLE_DOCUMENT = local_settings.SAMPLE_DOCUMENT


class Alignment:
    LEFT = 'left'
    RIGHT = 'right'


class HTMLElement:
    style = None
    text = None
    tag = None
    sub_tag = None
    children = None

    def __init__(self, tag, sub_tag=None, style=None, text=""):
        self.style = style
        self.sub_tag = sub_tag
        self.text = text
        self.tag = tag
        self.children = []


class HTMLParagraph:
    css = None
    text = None
    tag = None
    sub_tags = None

    def __init__(self, paragraph):

        bstyle_name = paragraph.style.base_style.name if paragraph.style.base_style else ""

        if bstyle_name == "List Paragraph":
            final_tag = 'li'
        else:
            final_tag = 'p'

        text_cells = paragraph.text.split('\t')
        tabbed_text_qty = len(text_cells)
        tabs = paragraph.paragraph_format.tab_stops
        tab_qty = len(tabs)
        low_tab_num = min(tabbed_text_qty, tab_qty)

        if low_tab_num > 1:
            self.text = ""
            self.css = "width: 100%;"
            self.tag = 'table'
            row = HTMLElement('tr')

            alignments = self.get_table_alignment(tabs)
            for t in range(low_tab_num):
                cell_css = self.build_css(paragraph, alignment_info=alignments[t])
                row.children.append(HTMLElement('td', sub_tag=final_tag,
                                                style=cell_css,
                                                text=text_cells[t]))
            self.sub_tags = [row]
        else:
            self.tag = final_tag
            self.text = paragraph.text
            self.css = self.build_css(paragraph)

    def build_css(self, paragraph, alignment_info=None):
        output_css = []

        text_align = "text-align: "

        if alignment_info:
            text_align += alignment_info
            output_css.append(text_align)

        font_size = "font-size:"

        if paragraph.style.font.size:
            font_size += "{}px".format(paragraph.style.font.size.pt)
        else:
            font_size += "{}px".format(DEFAULT_FONT_PT_SIZE)

        output_css.append(font_size)

        font_weight = "font-weight:"

        if paragraph.style.font.bold:
            font_weight += "bold"
            output_css.append(font_weight)

        font_style = "font-style:"

        if paragraph.style.font.italic:
            font_style += "italic"
            output_css.append(font_style)

        return ";".join(output_css)

    def get_table_alignment(self, tabs):
        if len(tabs) == 2 and \
                (tabs[0].alignment._member_name != 'RIGHT'
                 and tabs[1].alignment._member_name == 'RIGHT'):
                return [Alignment.LEFT, Alignment.RIGHT]
        else:
            return [Alignment.LEFT]*len(tabs)


class HTMLDocument:
    paragraphs = None

    def __init__(self, document):
        self.paragraphs = self.build_html_paragraphs(document)

    def build_html_paragraphs(self, document):
        output_paragraphs = []
        for p in document.paragraphs:
            output_paragraphs.append(HTMLParagraph(p))
        return output_paragraphs



def get_sample():
    return get_doc(SAMPLE_DOCUMENT)


def get_doc(filepath):
    docs = HTMLDocument(Document(filepath))
    return docs.paragraphs

