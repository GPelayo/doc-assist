from docx import Document
import re
import string
import os


TEXT_CHANGES = ""

try:
    import local_settings
except ImportError:
    SAMPLE_DOCUMENT = os.path.join("documents", "doc_templates", "sample.docx")
else:
    SAMPLE_DOCUMENT = local_settings.SAMPLE_DOCUMENT
    TEXT_CHANGES = local_settings.CHANGES


class Changes:
    paragraph_index = None
    target_phrase = None
    new_phrase = None
    token_offset = None
    paragraph_text = None
    original_case = None

    def __init__(self, paragraph, token, target_phrase, new_phrase, text, preserve_case=False):
        self.paragraph_text = text
        self.paragraph_index = paragraph
        self.token_offset = token
        self.new_phrase = new_phrase if preserve_case else match_case(text.split()[self.token_offset], new_phrase)
        self.target_phrase = target_phrase.lower()

    def __str__(self):
        return self.paragraph_text


def create_changes_for_phrase(target_phrase, new_phrase, preserve_case=False):
    docs = Document(SAMPLE_DOCUMENT)
    locations = []

    for p_i, p in enumerate(docs.paragraphs):
        locations += [Changes(p_i, t_i, target_phrase, new_phrase, p.text, preserve_case=preserve_case)
                      for t_i, t in enumerate(p.text.split()) if clean_word(t.lower()) == target_phrase.lower()]

    return locations


def match_case(original_text, new_text):
    fixed_text = clean_word(original_text)
    if fixed_text.islower():
        return new_text.lower()
    if fixed_text.isupper():
        return new_text.upper()
    if fixed_text.istitle():
        return new_text.title()
    else:
        return new_text



def clean_word(word):
    return re.sub('[' + string.punctuation + ']', '', word)


def create_new_doc(location_list):
    docs = Document(SAMPLE_DOCUMENT)
    for l in location_list:
        p = docs.paragraphs[l.paragraph_index]
        t_list = p.text.split(' ')
        t_list[l.token_offset] = t_list[l.token_offset].lower().replace(l.target_phrase, l.new_phrase)
        docs.paragraphs[l.paragraph_index].text = ' '.join(t_list)
    docs.save(os.path.join('output', 'new.docx'))



if __name__ == "__main__":
    change_list = []
    for c in TEXT_CHANGES:
        preserve_case = c.get('preserve_case', False)
        change_list += create_changes_for_phrase(c['target'], c['new'], preserve_case)
    create_new_doc(change_list)